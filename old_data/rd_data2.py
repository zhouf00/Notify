import xlrd, os, time
import pythoncom
from docx import Document
from win32com import client
from time import ctime, strftime, localtime



class ReadData(object):

    log_conut = 0
    ######################################
    #  数据初始化
    ######################################
    def __init__(self, *args, **kwargs):
        xl_name, xl_tb_name, doc_name = args
        self.file_path = os.path.join(os.getcwd())
        self.xl_path = os.path.join(self.file_path, xl_name)
        self.xl_tb_name = xl_tb_name
        self.doc_model_path = os.path.join(self.file_path, doc_name)
        log_dir = os.path.join(self.file_path,'log')
        if not os.path.exists(log_dir):
            os.mkdir(log_dir)
        self.log_path = os.path.join\
            (log_dir, strftime("%m%dlog.txt", localtime()))
        self.log_err = os.path.join \
            (log_dir, strftime("%m%dlog_all.txt", localtime()))
        self.init_xl()
    ######################################
    #  获取表格名、列、行
    ######################################
    def init_xl(self):
        data = xlrd.open_workbook(self.xl_path)
        table_name = data.sheet_names()
        name = [res for res in table_name if res.strip() == self.xl_tb_name]
        self.__xl_table = data.sheet_by_name(name[0])
        self.xl_rows = self.__xl_table.nrows
        self.xl_cols = self.__xl_table.ncols

    ######################################
    #  获取表格数据并存入迭代器
    ######################################
    def rd_xl_value(self):
        """
        接受行号，读取行信息
        :return:
        """
        for i in range(1, self.xl_rows):
            row_value = self.__xl_table.row_values(i,start_colx=0, end_colx=self.xl_cols-1)
            row_name = self.__xl_table.cell(i, 0).value
            row_email = self.__xl_table.cell(i, self.xl_cols-1).value
            yield row_name, row_email, row_value

    ######################################
    #  把每列数据写入word表格并保存
    ######################################
    def rd_doc(self, row_name, row_value):
        temp_dir = os.path.join(self.file_path, 'temp')
        if not os.path.exists(temp_dir):
            os.mkdir(temp_dir)
        doc_new_file = os.path.join(temp_dir,"%s.docx"%row_name)
        doc = Document(self.doc_model_path)
        self.doc_table = doc.tables[0]
        for i in range(len(row_value)):
            run = self.doc_table.cell(1, i).paragraphs[0].add_run(str(row_value[i]))
            run.font.size = 100000
        doc.save(doc_new_file)
        return doc_new_file

    ######################################
    #  word转为PDF文件，并删除word
    ######################################
    def doc2pdf(self, row_name, doc_new_file):
        pdf_file = os.path.join(self.file_path, "temp\%s.pdf" % row_name)
        pdf_name = os.path.basename(pdf_file)
        try:
            word = client.DispatchEx("word.Application")
            if os.path.exists(pdf_file):
                os.remove(pdf_file)
            self.worddoc = word.Documents.Open(doc_new_file, ReadOnly=1)
            #print("开始转换", ctime())
            self.worddoc.SaveAs(pdf_file, FileFormat=17)
            self.worddoc.Close()
            word.Quit() # 多线程使用时太费时间
            del(word)
            return pdf_name, pdf_file
        except Exception as e:
            pdferr = ("<%>PDF转换失败")%e
            #print(pdferr)
        finally:
            #print("PDF完成时间<%s>"%ctime())
            os.remove(doc_new_file)

    ######################################
    #  打印时间
    ######################################
    def p_time(self):
        return strftime("<%m-%d %H:%M:%S>", localtime())

    ######################################
    #  打开PDF的文件夹
    ######################################
    def _open_file(self):
        os.system("explorer %s"% os.path.join(self.file_path, "temp"))
        os.system("explorer %s" % self.log_path)

    ######################################
    #  打开PDF的文件夹，与发送前人员邮箱日志
    ######################################
    def _open_logall(self):
        os.system("explorer %s" % self.log_err)

    ######################################
    #  写入日志文件
    ######################################
    def _log_file(self, row_name, row_email, pdf_file, pdf_name):
        if self.log_conut == 0:
            f = open(self.log_path, "w", encoding="utf8")
            f.write(("%s,%s,%s,%s\n") % (row_name, row_email, pdf_name, pdf_file))
            self.log_conut += 1
        else:
            f = open(self.log_path, "a+", encoding="utf8")
            f.write(("%s,%s,%s,%s\n")%(row_name, row_email, pdf_name, pdf_file))

    ######################################
    #  实施存word，转PDF函数，并把完成信息写入列表
    ######################################
    def _change(self, full_value):
        pythoncom.CoInitialize()
        row_name, row_email, row_value = full_value
        doc_new_file = self.rd_doc(row_name, row_value)
        pdf_name, pdf_file = self.doc2pdf(row_name, doc_new_file)
        #self.a_list.append((row_name, row_email, pdf_file, pdf_name))
        return row_name, row_email, pdf_file, pdf_name

    def _rd_log(self):
        f = open(self.log_path, "r", encoding="utf8")
        #print(f.readlines())
        for var in f.readlines():
            yield [m_str.strip() for m_str in var.split(",")]
        #row_name, row_email, pdf_name, pdf_file = \
        #    [var.split(",")[1].strip() for var in f.readline()]
        f.close()

    def _err_log(self, var):
        if self.log_conut == 0:
            f = open(self.log_err, "w", encoding="utf8")
            f.write(var)
            self.log_conut += 1
        else:
            f = open(self.log_err, "a+", encoding="utf8")
            f.write(var)


if __name__ == '__main__':


    email_host = r"smtp.exmail.qq.com"
    email_port = 465
    mail_name = r"doc@windit.com.cn"
    mail_passwd = "Zzqa2018"
    name = r"文档管理"
    month = r"11月"
    sub = r"工资条"
    html_sign = r"D:\Study\项目练习\发邮件软件\qianmin.html"
    img_path = "D:\Study\项目练习\发邮件软件\公司logo.jpg"

    aa = ReadData("收款明细表样表.xls", "底稿模板", "工资条模板.docx")
    full_value = aa.rd_xl_value()

    try:
        while True:
            print(full_value.__next__())
    except StopIteration:
        exit()
    finally:
        print("测试成功", ctime())

