import os
import pandas as pd
import pythoncom
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from win32com import client


class FileProcessing(object):

    def __init__(self, data, tmpPath, doc_model):
        self._data_file = data
        self._doc_model = doc_model
        self._TEMP_PATH = tmpPath

    @property
    def get_data(self):
        df = pd.read_excel(self._data_file)
        df = df.where(df.notnull(), '')
        self._df = pd.DataFrame(df)
        self._columns = list(self._df)  # 获取列名
        self._columns.append('')
        return self._columns

    def get_excelData(self, s_str, e_str, type, to, cc):
        columns = self.get_data
        print(columns)
        # 截取有效数据
        # data = df.iloc[:, columns.index(s_str):columns.index(e_str) + 1]
        data = pd.concat([self._df.iloc[:, columns.index(s_str):columns.index(e_str) + 1],
                          self._df[to], self._df[type], self._df[cc]], axis=1)
        columns = list(data)
        # 获取收件人下标
        to_index = columns.index(to)
        # 获取信息归属下标
        type_index = columns.index(type)
        # 获取抄送人，无抄送人则为空
        if cc in columns:
            cc_index = columns.index(cc)
        else:
            cc_index = None
        data_dict = dict()
        # 将数据转化为字典格式
        # 用户 : {to: 发件人，cc: 抄送人，value: [[应发信息1],[应发信息2], ...]}
        for var in data.values:
            if var[to_index] in data_dict.keys():
                data_dict[var[type_index]]['value'].append(list(var[:to_index]))
            else:
                if cc_index:
                    data_dict[var[type_index]] = {
                        'to_user': var[to_index],
                        'cc_user': var[cc_index],
                        'value': [list(var[:columns.index(e_str)+1])]
                    }
                else:
                    data_dict[var[type_index]] = {
                        'to_user': var[to_index],
                        'cc_user': '',
                        'value': [list(var[:columns.index(e_str)+1])]
                    }

        return data_dict.items()

    ######################################
    #  进行文件转换的操作
    ######################################
    def spanned_file(self, name, data):

        pythoncom.CoInitialize()
        doc_file = self._rd_doc(name, data['value'])
        pdf_file = self._doc2pdf(name, doc_file)
        data['file'] = pdf_file # 添加PDF文件名到字典
        data.pop('value')   # 删除值
        return data

    ######################################
    #  把每列数据写入word表格并保存
    ######################################
    def _rd_doc(self, name, values):
        doc_new_file = os.path.join(self._TEMP_PATH, "%s.docx" % name)
        # print(self._doc_model)
        doc = Document(self._doc_model)
        doc_table = doc.tables[0]
        for value, col in zip(values, range(len(values))):
            # print(value)
            if col > 0:
                doc_table.add_row().cells
            for i in range(len(value)):
                # print('<%s>'%value[i],end=' ')
                run = doc_table.cell(col + 1, i).paragraphs[0]
                run.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                print(isinstance(value[i], str))
                if isinstance(value[i], str) or isinstance(value[i], int):
                    run = run.add_run(str(value[i]))
                else:
                    run = run.add_run(str(value[i].strftime("%Y-%m-%d")))

                run.font.size = 100000
            # print()
        doc.save(doc_new_file)
        return doc_new_file

    ######################################
    #  word转为PDF文件，并删除word
    ######################################
    def _doc2pdf(self, row_name, doc_new_file):
        pdf_file = os.path.join(self._TEMP_PATH, "%s.pdf" % row_name)
        # pdf_name = os.path.basename(pdf_file)
        try:
            word = client.DispatchEx("word.Application")
            if os.path.exists(pdf_file):
                os.remove(pdf_file)
            worddoc = word.Documents.Open(doc_new_file, ReadOnly=1)
            worddoc.SaveAs(pdf_file, FileFormat=17)
            worddoc.Close()
            word.Quit()  # 多线程使用时太费时间
            del (word)
            return pdf_file
        except Exception as e:
            pdferr = ("<%s>PDF转换失败") % e
            # print(pdferr)
        finally:
            pass
            # print("PDF完成时间<%s>"%ctime())
            os.remove(doc_new_file)


if __name__ == '__main__':
    init_path = os.path.join(os.getcwd())  # 初始路径
    temp_path = os.path.join(init_path, 'tmp')  # 文件转换路径
    if not os.path.exists(temp_path):
        os.mkdir(temp_path)
    email_file = os.path.join(init_path, '邮箱管理.xlsx')
    data_file = os.path.join(init_path, '付款明细表样表.xls')
    doc_model_file = os.path.join(init_path, '模板', '付款明细模板.docx')
    S = '日期'
    E = '金额'
    TO = '报销人'
    CC = None
    d = FileProcessing(data_file, temp_path)

    datas = d.get_excelData(S, E, TO, TO, CC)
    for name, data in datas:
        d.spanned_file(name, data)
        pass
    emails = dict(pd.DataFrame(pd.read_excel(email_file),columns=['名字','邮箱']).values)
    print(emails.get('周凡'))

